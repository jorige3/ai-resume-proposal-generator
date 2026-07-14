from io import BytesIO
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor

from app.models.schemas import UserProfile, GeneratedResumeContent, GeneratedProposal, JobMatchScore
from app.services.pdf_exporter import ExportError


def _set_font(run, name='Arial', size_pt=10.5, color_rgb=(45, 55, 72), bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.bold = bold
    run.italic = italic


def _add_styled_paragraph(doc, text, is_bullet=False, bold_prefix=None, space_after=6, font_size=10.5, color_rgb=(45, 55, 72), italic=False):
    style = 'List Bullet' if is_bullet else 'Normal'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        _set_font(r_pre, size_pt=font_size, color_rgb=color_rgb, bold=True)
        
    r_text = p.add_run(text)
    _set_font(r_text, size_pt=font_size, color_rgb=color_rgb, italic=italic)
    return p


def _add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    _set_font(run, size_pt=13, color_rgb=(30, 58, 138), bold=True)
    
    # Add a thin line under the section heading by creating a paragraph with a border or using a separator line
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_after = Pt(6)
    p_sep.paragraph_format.space_before = Pt(0)
    p_sep.paragraph_format.line_spacing = 1.0
    r_sep = p_sep.add_run("―" * 50)
    _set_font(r_sep, size_pt=8, color_rgb=(226, 232, 240))


def export_resume_docx(
    profile: UserProfile,
    resume: GeneratedResumeContent,
    match_score: Optional[JobMatchScore] = None
) -> bytes:
    """Generates a professional resume Word document in memory and returns its bytes."""
    if not profile.full_name or not profile.full_name.strip():
        raise ExportError("Candidate full name cannot be empty.")
    if not resume.professional_summary or not resume.professional_summary.strip():
        raise ExportError("Resume professional summary cannot be empty.")
        
    try:
        doc = Document()
        
        # Page Margins: 0.75 inch (54pt)
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            
        # 1. Candidate Name
        p_name = doc.add_paragraph()
        p_name.paragraph_format.space_after = Pt(2)
        r_name = p_name.add_run(profile.full_name)
        _set_font(r_name, size_pt=20, color_rgb=(30, 58, 138), bold=True)
        
        # Title
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_after = Pt(6)
        r_title = p_title.add_run(profile.title)
        _set_font(r_title, size_pt=12, color_rgb=(74, 85, 104), italic=True)
        
        # 2. Contact details
        contact_parts = []
        if profile.email and profile.email.strip():
            contact_parts.append(f"Email: {profile.email.strip()}")
        if profile.phone and profile.phone.strip():
            contact_parts.append(f"Phone: {profile.phone.strip()}")
        if contact_parts:
            p_contact = doc.add_paragraph()
            p_contact.paragraph_format.space_after = Pt(12)
            r_contact = p_contact.add_run(" | ".join(contact_parts))
            _set_font(r_contact, size_pt=9.5, color_rgb=(113, 128, 150))
            
        # 3. Job Match Summary Table
        if match_score:
            table = doc.add_table(rows=2, cols=3)
            table.autofit = True
            
            # Row 0
            r0 = table.rows[0]
            _set_font(r0.cells[0].paragraphs[0].add_run(f"Match Score: {match_score.overall_score}%"), size_pt=9.5, color_rgb=(30, 58, 138), bold=True)
            _set_font(r0.cells[1].paragraphs[0].add_run(f"Skills Fit: {match_score.skills_match_score}%"), size_pt=9.5, color_rgb=(30, 58, 138), bold=True)
            _set_font(r0.cells[2].paragraphs[0].add_run(f"Experience Fit: {match_score.experience_match_score}%"), size_pt=9.5, color_rgb=(30, 58, 138), bold=True)
            
            # Row 1
            r1 = table.rows[1]
            merged = r1.cells[0].merge(r1.cells[1]).merge(r1.cells[2])
            p_exp = merged.paragraphs[0]
            _set_font(p_exp.add_run(f"Overview: {match_score.explanation}"), size_pt=9.5, color_rgb=(74, 85, 104), italic=True)
            
            # Formatting table spacer
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_after = Pt(8)
            
        # 4. Professional Summary
        _add_section_heading(doc, "Professional Summary")
        _add_styled_paragraph(doc, resume.professional_summary)
        
        # 5. Key Skills
        if resume.tailored_skills:
            _add_section_heading(doc, "Key Skills")
            _add_styled_paragraph(doc, ", ".join(resume.tailored_skills))
            
        # 6. Work Experience
        if profile.experience:
            _add_section_heading(doc, "Professional Experience")
            for exp in profile.experience:
                end = exp.end_date if exp.end_date else "Present"
                p_exp_header = doc.add_paragraph()
                p_exp_header.paragraph_format.space_after = Pt(2)
                p_exp_header.paragraph_format.keep_with_next = True
                
                r_role_comp = p_exp_header.add_run(f"{exp.role} at {exp.company}")
                _set_font(r_role_comp, size_pt=11, color_rgb=(30, 58, 138), bold=True)
                
                r_dates = p_exp_header.add_run(f"  ({exp.start_date} - {end})")
                _set_font(r_dates, size_pt=9.5, color_rgb=(113, 128, 150))
                
                _add_styled_paragraph(doc, exp.description, space_after=8)
                
        # 7. Projects
        if profile.projects:
            _add_section_heading(doc, "Projects")
            for proj in profile.projects:
                p_proj_header = doc.add_paragraph()
                p_proj_header.paragraph_format.space_after = Pt(2)
                p_proj_header.paragraph_format.keep_with_next = True
                
                r_title = p_proj_header.add_run(proj.title)
                _set_font(r_title, size_pt=11, color_rgb=(30, 58, 138), bold=True)
                
                if proj.url and proj.url.strip():
                    r_url = p_proj_header.add_run(f"  ({proj.url.strip()})")
                    _set_font(r_url, size_pt=9, color_rgb=(113, 128, 150))
                    
                if proj.technologies:
                    p_tech = doc.add_paragraph()
                    p_tech.paragraph_format.space_after = Pt(2)
                    p_tech.paragraph_format.keep_with_next = True
                    r_tech = p_tech.add_run(f"Technologies: {', '.join(proj.technologies)}")
                    _set_font(r_tech, size_pt=9.5, color_rgb=(74, 85, 104), italic=True)
                    
                _add_styled_paragraph(doc, proj.description, space_after=8)
                
        # 8. Education
        if profile.education:
            _add_section_heading(doc, "Education")
            for edu in profile.education:
                grad_yr = f"  (Graduated: {edu.graduation_year})" if edu.graduation_year else ""
                p_edu = doc.add_paragraph()
                p_edu.paragraph_format.space_after = Pt(4)
                
                r_deg = p_edu.add_run(f"{edu.degree} in {edu.field_of_study}")
                _set_font(r_deg, size_pt=10.5, color_rgb=(30, 58, 138), bold=True)
                
                r_inst = p_edu.add_run(f" at {edu.institution}{grad_yr}")
                _set_font(r_inst, size_pt=10.5, color_rgb=(45, 55, 72))
                
        # 9. Certifications
        if profile.certifications:
            _add_section_heading(doc, "Certifications")
            for cert in profile.certifications:
                yr_str = f" ({cert.year})" if cert.year else ""
                cert_desc = f"{cert.name} issued by {cert.issuing_org}{yr_str}"
                _add_styled_paragraph(doc, cert_desc, is_bullet=True, space_after=3)
                
        # 10. Suggestions
        if resume.improvement_suggestions or resume.missing_skills_report:
            doc.add_page_break()
            _add_section_heading(doc, "Resume Tailoring Insights & Recommendations")
            
            if resume.improvement_suggestions:
                p_sug_title = doc.add_paragraph()
                p_sug_title.paragraph_format.space_after = Pt(4)
                r_sug_title = p_sug_title.add_run("Improvement Suggestions:")
                _set_font(r_sug_title, size_pt=11, color_rgb=(30, 58, 138), bold=True)
                for sug in resume.improvement_suggestions:
                    _add_styled_paragraph(doc, sug, is_bullet=True, space_after=3)
                    
            if resume.missing_skills_report:
                p_miss_title = doc.add_paragraph()
                p_miss_title.paragraph_format.space_before = Pt(8)
                p_miss_title.paragraph_format.space_after = Pt(4)
                r_miss_title = p_miss_title.add_run("Missing Skills & Recommendations:")
                _set_font(r_miss_title, size_pt=11, color_rgb=(30, 58, 138), bold=True)
                for rep in resume.missing_skills_report:
                    _add_styled_paragraph(doc, rep, is_bullet=True, space_after=3)
                    
        # Write doc to buffer
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    except Exception as e:
        raise ExportError(f"Failed to generate resume Word document: {str(e)}") from e


def export_proposal_docx(
    profile: UserProfile,
    proposal: GeneratedProposal,
    match_score: Optional[JobMatchScore] = None
) -> bytes:
    """Generates a professional freelance proposal Word document in memory and returns its bytes."""
    if not profile.full_name or not profile.full_name.strip():
        raise ExportError("Candidate full name cannot be empty.")
    if not proposal.proposal_text or not proposal.proposal_text.strip():
        raise ExportError("Proposal text content cannot be empty.")
        
    try:
        doc = Document()
        
        # Margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            
        # 1. Candidate Name
        p_name = doc.add_paragraph()
        p_name.paragraph_format.space_after = Pt(2)
        r_name = p_name.add_run(profile.full_name)
        _set_font(r_name, size_pt=20, color_rgb=(30, 58, 138), bold=True)
        
        # Title
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_after = Pt(6)
        r_title = p_title.add_run(profile.title)
        _set_font(r_title, size_pt=12, color_rgb=(74, 85, 104), italic=True)
        
        # 2. Contact details
        contact_parts = []
        if profile.email and profile.email.strip():
            contact_parts.append(f"Email: {profile.email.strip()}")
        if profile.phone and profile.phone.strip():
            contact_parts.append(f"Phone: {profile.phone.strip()}")
        if contact_parts:
            p_contact = doc.add_paragraph()
            p_contact.paragraph_format.space_after = Pt(12)
            r_contact = p_contact.add_run(" | ".join(contact_parts))
            _set_font(r_contact, size_pt=9.5, color_rgb=(113, 128, 150))
            
        # 3. Compatibility score
        if match_score:
            table = doc.add_table(rows=1, cols=3)
            table.autofit = True
            r0 = table.rows[0]
            _set_font(r0.cells[0].paragraphs[0].add_run(f"Overall Match: {match_score.overall_score}%"), size_pt=9.5, color_rgb=(30, 58, 138), bold=True)
            _set_font(r0.cells[1].paragraphs[0].add_run(f"Skills Fit: {match_score.skills_match_score}%"), size_pt=9.5, color_rgb=(30, 58, 138), bold=True)
            _set_font(r0.cells[2].paragraphs[0].add_run(f"Experience Fit: {match_score.experience_match_score}%"), size_pt=9.5, color_rgb=(30, 58, 138), bold=True)
            
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_after = Pt(8)
            
        # 4. Proposal Heading & Body
        _add_section_heading(doc, "Freelance Proposal")
        _add_styled_paragraph(doc, proposal.proposal_text)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    except Exception as e:
        raise ExportError(f"Failed to generate proposal Word document: {str(e)}") from e
